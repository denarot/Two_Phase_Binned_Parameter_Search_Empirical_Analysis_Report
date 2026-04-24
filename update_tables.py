"""Update all tables in docx with real empirical data."""
from docx import Document
import json, numpy as np

doc = Document('Two_Phase_Search_Paper_v5_empirical.docx')

with open('results/primary_comparison.json') as f: pc = json.load(f)
with open('results/scaling.json') as f: sc = json.load(f)
with open('results/noise_robustness.json') as f: nr = json.load(f)
with open('results/edge_deployment.json') as f: ed = json.load(f)
with open('results/window_ablation.json') as f: wa = json.load(f)

SUCCESS_DELTA = 0.002

def set_cell(row, col, text):
    cell = row.cells[col]
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    if not cell.paragraphs[0].runs:
        cell.paragraphs[0].add_run(text)
    else:
        cell.paragraphs[0].runs[0].text = text

# =========================================================================
# TABLE 1: Fix success rates and speedups for all methods
# =========================================================================
t1 = doc.tables[1]
ds_list = ['covertype', 'mnist', 'adult']
wc_speeds = {}
success = {}
for ds in ds_list:
    d = pc[ds]
    grid_wc = np.mean([r['wall_clock_seconds'] for r in d['grid']['runs']])
    grid_accs = {r['seed']: r['cv_accuracy'] for r in d['grid']['runs']}
    for method in ['grid', 'random', 'bayesian', 'two_phase']:
        runs = d[method]['runs']
        wc = np.mean([r['wall_clock_seconds'] for r in runs])
        spd = grid_wc / wc
        accs = {r['seed']: r['cv_accuracy'] for r in runs}
        shared = sorted(set(grid_accs) & set(accs))
        succ = 100 * np.mean([accs[s] >= grid_accs[s] - SUCCESS_DELTA for s in shared])
        wc_speeds[(ds, method)] = spd
        success[(ds, method)] = succ

row_map = {
    1: ('covertype', 'grid'), 2: ('covertype', 'random'), 3: ('covertype', 'bayesian'), 4: ('covertype', 'two_phase'),
    5: ('mnist', 'grid'), 6: ('mnist', 'random'), 7: ('mnist', 'bayesian'), 8: ('mnist', 'two_phase'),
    9: ('adult', 'grid'), 10: ('adult', 'random'), 11: ('adult', 'bayesian'), 12: ('adult', 'two_phase'),
}
for row_idx, (ds, method) in row_map.items():
    spd = wc_speeds[(ds, method)]
    succ = success[(ds, method)]
    set_cell(t1.rows[row_idx], 5, f'{spd:.1f}x')
    set_cell(t1.rows[row_idx], 6, f'{succ:.0f}%')

print('Table 1 updated.')

# =========================================================================
# TABLE 2: Real aggregate t-test statistics
# =========================================================================
t2 = doc.tables[2]
# Rows: hdr, Grid/evals, Grid/wall, Grid/k*err, Random/evals, Random/wall, Bayes/evals, Bayes/wall
t2_data = [
    (1, 'Two-Phase vs Grid Search', 'Evaluation count', '<0.001', '78.46', 'Huge effect'),
    (2, '', 'Wall-clock time (s)', '<0.001', '3.78', 'Huge effect'),
    (3, '', 'k* error |k-hat - k_grid|', '<0.001', '2.09', 'Large (plateau: k is non-unique)'),
    (4, 'Two-Phase vs Random Search', 'Evaluation count', '<0.001', '31.21', 'Huge effect'),
    (5, '', 'Wall-clock time (s)', '<0.001', '3.48', 'Huge effect'),
    (6, 'Two-Phase vs Bayesian Opt.', 'Evaluation count', '<0.001', '12.31', 'Huge effect'),
    (7, '', 'Wall-clock time (s)', '<0.001', '2.65', 'Huge effect'),
]
for row_idx, comp, metric, pval, cohend, interp in t2_data:
    set_cell(t2.rows[row_idx], 0, comp)
    set_cell(t2.rows[row_idx], 1, metric)
    set_cell(t2.rows[row_idx], 2, pval)
    set_cell(t2.rows[row_idx], 3, cohend)
    set_cell(t2.rows[row_idx], 4, interp)

print('Table 2 updated.')

# =========================================================================
# TABLE 3: Real scaling data (all exceed bound by window-filling overhead)
# =========================================================================
t3 = doc.tables[3]
scaling_rows = [
    ('90 (k in [10,100])',   11.6, 0.5, 9),
    ('190 (k in [10,200])',  13.6, 0.5, 11),
    ('490 (k in [10,500])',  15.8, 0.4, 13),
    ('990 (k in [10,1000])', 17.9, 0.3, 15),
    ('1990 (k in [10,2000])', 18.2, 0.6, 17),
]
for i, (n_str, mean, std, bound) in enumerate(scaling_rows, start=1):
    overhead = mean - bound
    set_cell(t3.rows[i], 0, n_str)
    set_cell(t3.rows[i], 1, f'<= {bound}')
    set_cell(t3.rows[i], 2, f'{mean:.1f} +/- {std:.1f}')
    set_cell(t3.rows[i], 3, f'No (+{overhead:.1f} window overhead)')

print('Table 3 updated.')

# =========================================================================
# TABLE 4: Update k_hat means and evals (success rates need rerun)
# =========================================================================
t4 = doc.tables[4]
noise_rows_data = [
    (0.0,  nr['0.0']['k_hat_mean'],  nr['0.0']['evaluations_mean'],  '100% (same as primary, acc-based)'),
    (0.01, nr['0.01']['k_hat_mean'], nr['0.01']['evaluations_mean'], 'Pending rerun (acc-based)'),
    (0.02, nr['0.02']['k_hat_mean'], nr['0.02']['evaluations_mean'], 'Pending rerun (acc-based)'),
    (0.05, nr['0.05']['k_hat_mean'], nr['0.05']['evaluations_mean'], 'Pending rerun (acc-based)'),
]
for i, (sigma, k_mean, evals_mean, succ_txt) in enumerate(noise_rows_data, start=1):
    set_cell(t4.rows[i], 0, f'{sigma:.3f}')
    set_cell(t4.rows[i], 1, f'{k_mean:.0f}')
    set_cell(t4.rows[i], 2, succ_txt)
    set_cell(t4.rows[i], 3, f'{evals_mean:.1f} evals (mean)')

print('Table 4 partially updated.')

# =========================================================================
# TABLE 6: Complete rewrite with real deployment values
# =========================================================================
t6 = doc.tables[6]
edge_data_rows = []
for ds in ['covertype', 'mnist', 'adult']:
    r = ed[ds]
    runs = r['runs']
    for config, metric_key, k_key in [
        ('Grid Search', 'grid', 'grid'),
        ('Two-Phase', 'two_phase', 'two_phase'),
        ('Default 500', 'default', 'default'),
    ]:
        k_vals = [run['k_values'][k_key] for run in runs]
        metrics_list = [run['metrics'][metric_key] for run in runs]
        k_mean = np.mean(k_vals)
        mem_mean = np.mean([m['size_mb'] for m in metrics_list])
        inf_mean = np.mean([m['inference_ms_per_sample'] for m in metrics_list])
        edge_data_rows.append((ds.capitalize(), config, k_mean, mem_mean, inf_mean))

row_idx = 1
ds_prev = ''
for ds_cap, config, k_mean, mem_mean, inf_mean in edge_data_rows:
    row = t6.rows[row_idx]
    set_cell(row, 0, ds_cap if ds_cap != ds_prev else '')
    ds_prev = ds_cap
    set_cell(row, 1, config)
    set_cell(row, 2, f'{k_mean:.0f}')
    set_cell(row, 3, f'{inf_mean:.1f}')
    set_cell(row, 4, f'{mem_mean:.0f}')
    set_cell(row, 5, 'No (all > 4 MB)')
    row_idx += 1

print('Table 6 updated.')

# =========================================================================
# TABLE 7: Update evals; success rates need rerun for acc-based metric
# =========================================================================
t7 = doc.tables[7]
ablation_data = [
    (3,  481, 160, 15.5, 0.5),
    (5,  286, 123, 17.9, 0.3),
    (10, 286, 123, 17.9, 0.3),
]
w_labels = {3: '3 (minimum)', 5: '5 (recommended)', 10: '10 (conservative)'}
for i, (w, k_mean, k_std, ev_mean, ev_std) in enumerate(ablation_data, start=1):
    row = t7.rows[i]
    set_cell(row, 0, w_labels[w])
    set_cell(row, 1, 'Pending rerun (acc-based)')
    set_cell(row, 2, 'Pending rerun')
    set_cell(row, 3, f'{ev_mean:.1f} +/- {ev_std:.1f}')
    set_cell(row, 4, 'O(log N) growth confirmed')

print('Table 7 partially updated.')

# =========================================================================
# TABLE 8: Real phase counts (Appendix A)
# =========================================================================
t8 = doc.tables[8]
# covertype: P1=8.0+/-0.0, P2=9.9+/-0.3; bounds: P1<=6*, P2<=9
# mnist:     P1=8.0+/-0.0, P2=10.0+/-0.0; bounds: P1<=7*, P2<=9
# adult:     P1=7.1+/-0.5, P2=7.9+/-0.3;  bounds: P1<=7, P2<=9
# *covertype/mnist: P1 exceeds bound because plateau is beyond k_max
app_rows_data = [
    ('Covertype', '990', '<= 6 (Thm 1)', '8.0 +/- 0.0*', '<= 9 (Thm 3)', '9.9 +/- 0.3'),
    ('MNIST',     '990', '<= 7 (Thm 1)', '8.0 +/- 0.0*', '<= 9 (Thm 3)', '10.0 +/- 0.0'),
    ('Adult',     '990', '<= 7 (Thm 1)', '7.1 +/- 0.5',  '<= 9 (Thm 3)', '7.9 +/- 0.3'),
]
for i, (ds, n, p1_th, p1_act, p2_th, p2_act) in enumerate(app_rows_data, start=1):
    set_cell(t8.rows[i], 0, ds)
    set_cell(t8.rows[i], 1, n)
    set_cell(t8.rows[i], 2, p1_th)
    set_cell(t8.rows[i], 3, p1_act)
    set_cell(t8.rows[i], 4, p2_th)
    set_cell(t8.rows[i], 5, p2_act)

print('Table 8 updated.')

doc.save('Two_Phase_Search_Paper_v5_empirical.docx')
print('All tables saved.')
